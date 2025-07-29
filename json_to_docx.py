#!/usr/bin/env python3
"""
JSON to DOCX Converter with Track Changes or Comments

This script takes JSON input containing content and comments, then creates a Word document
with the content and inserts either track changes or comments at the specified start and end positions.
"""

import json
import sys
import argparse
from typing import List, Dict, Any
from docx import Document
from docx.shared import Inches
from docx.oxml.shared import OxmlElement, qn
from docx.oxml.ns import nsdecls
from docx.oxml import parse_xml
from bs4 import BeautifulSoup
import re
from datetime import datetime


def create_revision_element(comment_text: str):
    """Create a track changes revision element."""
    # Create the revision element
    revision = OxmlElement('w:del')
    revision.set(qn('w:author'), 'Reviewer')
    revision.set(qn('w:date'), datetime.now().strftime('%Y-%m-%dT%H:%M:%SZ'))
    
    # Add the deleted text
    text_element = OxmlElement('w:t')
    text_element.text = comment_text
    revision.append(text_element)
    
    return revision


def add_comment_to_paragraph(paragraph, document: Document, start_index: int, end_index: int, comment_id: str, comment_text: str):
    """Add a comment to a paragraph at the specified text range."""
    # Get the text of the paragraph
    text = paragraph.text
    
    if start_index >= len(text) or end_index > len(text) or start_index >= end_index:
        print(f"Warning: Invalid comment range {start_index}-{end_index} for text length {len(text)}")
        return
    
    # Clear the paragraph
    paragraph.clear()
    
    # Split text into parts: before comment, commented text, after comment
    before_comment = text[:start_index]
    commented_text = text[start_index:end_index]
    after_comment = text[end_index:]
    
    # Add text before comment
    if before_comment:
        paragraph.add_run(before_comment)
    
    # Add commented text
    commented_run = paragraph.add_run(commented_text)
    
    # Add text after comment
    if after_comment:
        paragraph.add_run(after_comment)
    
    # Add comment using the proper python-docx method
    try:
        # Get the run that contains the commented text
        runs = paragraph.runs
        # Find the run that contains our commented text
        target_run = None
        for run in runs:
            if commented_text in run.text:
                target_run = run
                break
        
        if target_run:
            # Add comment to the specific run
            document.add_comment(
                runs=[target_run],
                text=comment_text,
                author="Reviewer",
                initials="R"
            )
        else:
            # Fallback: add comment to the paragraph
            document.add_comment(
                runs=runs,
                text=comment_text,
                author="Reviewer",
                initials="R"
            )
    except Exception as e:
        print(f"Warning: Could not add comment: {e}")


def add_track_change_to_paragraph(paragraph, start_index: int, end_index: int, comment_text: str):
    """Add a track change to a paragraph at the specified text range."""
    # Get the text of the paragraph
    text = paragraph.text
    
    if start_index >= len(text) or end_index > len(text) or start_index >= end_index:
        print(f"Warning: Invalid track change range {start_index}-{end_index} for text length {len(text)}")
        return
    
    # Clear the paragraph
    paragraph.clear()
    
    # Split text into parts: before change, changed text, after change
    before_change = text[:start_index]
    changed_text = text[start_index:end_index]
    after_change = text[end_index:]
    
    # Add text before change
    if before_change:
        paragraph.add_run(before_change)
    
    # Add track change (deletion) for the original text
    revision = create_revision_element(changed_text)
    
    # Add the revision as a deletion
    run = paragraph.add_run()
    run._element.append(revision)
    
    # Add the replacement text (this will appear as an insertion)
    replacement_run = paragraph.add_run(comment_text)
    # Mark as insertion
    insertion = OxmlElement('w:ins')
    insertion.set(qn('w:author'), 'Reviewer')
    insertion.set(qn('w:date'), datetime.now().strftime('%Y-%m-%dT%H:%M:%SZ'))
    replacement_run._element.getparent().replace(replacement_run._element, insertion)
    insertion.append(replacement_run._element)
    
    # Add text after change
    if after_change:
        paragraph.add_run(after_change)


def html_to_docx_paragraphs(document: Document, html_content: str, comments: List[Dict] = None, use_track_changes: bool = True):
    """Convert HTML content to Word document paragraphs with comments or track changes."""
    if not html_content:
        return
    
    # Parse HTML
    soup = BeautifulSoup(html_content, 'html.parser')
    
    # Process each HTML element
    for element in soup.find_all(['p', 'h1', 'h2', 'h3', 'h4', 'h5', 'h6', 'ul', 'ol', 'li']):
        if element.name in ['h1', 'h2', 'h3', 'h4', 'h5', 'h6']:
            # Add heading
            level = int(element.name[1])
            paragraph = document.add_heading(element.get_text(), level=level)
        elif element.name == 'ul':
            # Add unordered list
            for li in element.find_all('li', recursive=False):
                paragraph = document.add_paragraph()
                paragraph.style = 'List Bullet'
                paragraph.add_run(li.get_text())
        elif element.name == 'ol':
            # Add ordered list
            for li in element.find_all('li', recursive=False):
                paragraph = document.add_paragraph()
                paragraph.style = 'List Number'
                paragraph.add_run(li.get_text())
        elif element.name == 'li':
            # Skip individual li elements as they're handled in ul/ol
            continue
        else:
            # Add regular paragraph
            paragraph = document.add_paragraph(element.get_text())
        
        # Add comments or track changes if they exist for this paragraph
        if comments:
            paragraph_text = paragraph.text
            for comment in comments:
                start_idx = comment.get('start_index', 0)
                end_idx = comment.get('end_index', 0)
                comment_content = comment.get('comment_content', '')
                comment_id = comment.get('id', '0')
                
                # Check if comment range is within this paragraph
                if start_idx < len(paragraph_text) and end_idx <= len(paragraph_text):
                    if use_track_changes:
                        add_track_change_to_paragraph(paragraph, start_idx, end_idx, comment_content)
                    else:
                        add_comment_to_paragraph(paragraph, document, start_idx, end_idx, comment_id, comment_content)
                    break  # Only handle one comment per paragraph for now


def process_json_to_docx(json_data: List[Dict], output_filename: str = "output.docx", use_track_changes: bool = True):
    """Process JSON data and create a Word document with comments or track changes."""
    document = Document()
    
    for item in json_data:
        field_name = item.get('field_name', '')
        content = item.get('content', '')
        content_type = item.get('content_type', 'text')
        comments = item.get('comments', [])
        
        # Add field name as a heading if it exists
        if field_name:
            document.add_heading(field_name.replace('_', ' ').title(), level=1)
        
        # Process content based on type
        if content_type.lower() == 'html':
            html_to_docx_paragraphs(document, content, comments, use_track_changes)
        else:
            # Handle plain text
            paragraph = document.add_paragraph(content)
            
            # Add comments or track changes if they exist
            if comments:
                for comment in comments:
                    start_idx = comment.get('start_index', 0)
                    end_idx = comment.get('end_index', 0)
                    comment_content = comment.get('comment_content', '')
                    comment_id = comment.get('id', '0')
                    
                    if use_track_changes:
                        add_track_change_to_paragraph(paragraph, start_idx, end_idx, comment_content)
                    else:
                        add_comment_to_paragraph(paragraph, document, start_idx, end_idx, comment_id, comment_content)
        
        # Add some spacing between sections
        document.add_paragraph()
    
    # Save the document
    document.save(output_filename)
    print(f"Document saved as {output_filename}")


def main():
    """Main function to handle command line usage."""
    parser = argparse.ArgumentParser(description='Convert JSON to DOCX with comments or track changes')
    parser.add_argument('input', help='Input JSON file or "-" for stdin')
    parser.add_argument('output', nargs='?', default='output.docx', help='Output DOCX file (default: output.docx)')
    parser.add_argument('--trackchanges', action='store_true', help='Use track changes (default)')
    parser.add_argument('--comments', action='store_true', help='Use comments instead of track changes')
    
    args = parser.parse_args()
    
    # Determine mode
    use_track_changes = True  # Default
    if args.comments:
        use_track_changes = False
    elif args.trackchanges:
        use_track_changes = True
    
    # Check if input is a file or piped content
    if args.input == '-':
        # Read from stdin
        json_content = sys.stdin.read()
    else:
        # Read from file
        try:
            with open(args.input, 'r', encoding='utf-8') as f:
                json_content = f.read()
        except FileNotFoundError:
            print(f"Error: File {args.input} not found")
            sys.exit(1)
        except Exception as e:
            print(f"Error reading file: {e}")
            sys.exit(1)
    
    # Parse JSON
    try:
        json_data = json.loads(json_content)
    except json.JSONDecodeError as e:
        print(f"Error parsing JSON: {e}")
        sys.exit(1)
    
    # Process the data
    try:
        mode = "track changes" if use_track_changes else "comments"
        print(f"Creating document with {mode}...")
        process_json_to_docx(json_data, args.output, use_track_changes)
        print("Document created successfully!")
    except Exception as e:
        print(f"Error creating document: {e}")
        sys.exit(1)


if __name__ == "__main__":
    main() 